import os
from docx import Document

def create_resume(filename, name, email, text_lines):
    doc = Document()
    doc.add_heading(name, 0)
    doc.add_paragraph(f"Email: {email}")
    doc.add_paragraph("---------------------------------------")
    for line in text_lines:
        if line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=1)
        elif line.startswith("- "):
            doc.add_paragraph(line, style='List Bullet')
        else:
            doc.add_paragraph(line)
            
    os.makedirs("../test_resumes", exist_ok=True)
    doc.save(os.path.join("../test_resumes", filename))
    print(f"Created {filename}")

if __name__ == "__main__":
    # 1. Alice - Python Developer (Strong Match)
    create_resume(
        "alice_python_developer.docx",
        "Alice Johnson",
        "alice.johnson@email.com",
        [
            "## Summary",
            "Highly skilled Software Engineer with 4 years of experience specializing in Python, backend development, and RESTful APIs.",
            "## Skills",
            "- Languages: Python, JavaScript, SQL",
            "- Frameworks: Flask, Django, FastAPI",
            "- DevOps & Tools: Docker, Kubernetes, Git, PostgreSQL, Redis",
            "- Core concepts: Microservices, REST APIs, CI/CD, Unit Testing",
            "## Experience",
            "Senior Backend Engineer at TechCorp (2023 - Present)",
            "- Designed and implemented scalable REST APIs using Flask and Django",
            "- Containerized applications using Docker and orchestrated services with Kubernetes",
            "- Optimized SQL database queries, improving performance by 40%",
            "Software Developer at WebSystems (2021 - 2023)",
            "- Developed backend microservices in Python using FastAPI",
            "- Participated in Agile development, sprint planning, and code reviews",
            "## Education",
            "B.S. in Computer Science - State University",
            "## Certifications",
            "- Certified Kubernetes Administrator (CKA)",
            "- AWS Certified Solutions Architect"
        ]
    )

    # 2. Bob - Data Analyst / Python Scripting (Partial Match)
    create_resume(
        "bob_data_analyst.docx",
        "Bob Smith",
        "bob.smith@email.com",
        [
            "## Summary",
            "Data Analyst with 2+ years of experience processing large datasets, generating reports, and writing Python automation scripts.",
            "## Skills",
            "- Languages: Python, SQL, R",
            "- Libraries: Pandas, NumPy, Matplotlib, Scikit-Learn",
            "- Tools: Tableau, Excel, Git, MySQL",
            "## Experience",
            "Data Analyst at Analytics Ltd (2022 - Present)",
            "- Created Python scripts to automate ETL pipelines, saving 15 hours of manual work per week",
            "- Built SQL queries to extract business KPIs and designed interactive Tableau dashboards",
            "- Utilized Pandas and NumPy for data cleaning and statistical analysis",
            "## Education",
            "B.S. in Mathematics and Statistics - City College"
        ]
    )

    # 3. Charlie - UI/UX Designer (Low Match)
    create_resume(
        "charlie_designer.docx",
        "Charlie Brown",
        "charlie.brown@email.com",
        [
            "## Summary",
            "Creative UI/UX Designer with a passion for designing clean, intuitive web and mobile interfaces. Experienced in wireframing, prototyping, and user research.",
            "## Skills",
            "- Design Tools: Figma, Sketch, Adobe XD, Photoshop, Illustrator",
            "- Frontend Tech: HTML5, CSS3, JavaScript, TailwindCSS, React.js",
            "- Methodologies: User Research, Wireframing, High-Fidelity Prototyping, Usability Testing",
            "## Experience",
            "UI/UX Designer at Creative Studio (2023 - Present)",
            "- Designed beautiful user journeys and wireframes for e-commerce websites and mobile apps",
            "- Created interactive prototypes in Figma for client presentations",
            "- Collaborated with frontend engineers using React and TailwindCSS to ensure design accuracy",
            "## Education",
            "Bachelor of Fine Arts in Graphic Design - Academy of Art"
        ]
    )
